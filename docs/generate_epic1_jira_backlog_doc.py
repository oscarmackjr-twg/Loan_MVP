from datetime import date

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def main() -> None:
    out_path = r"c:\Users\omack\Intrepid\pythonFramework\cursor_loan_engine\docs\Epic1_Jira_Backlog.docx"
    doc = Document()

    p = doc.add_paragraph()
    r = p.add_run("Epic 1 Jira-Ready Backlog")
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run("Daily Loan Analysis & Wire Instructions\n").italic = True
    p.add_run(f"Prepared: {date.today().isoformat()}")

    doc.add_paragraph()
    doc.add_heading("Epic", level=1)
    epic_table = doc.add_table(rows=1, cols=6)
    hdr = epic_table.rows[0].cells
    hdr[0].text = "Issue Type"
    hdr[1].text = "Issue Key"
    hdr[2].text = "Summary"
    hdr[3].text = "Outcome"
    hdr[4].text = "Story Points"
    hdr[5].text = "Dependencies"
    row = epic_table.add_row().cells
    row[0].text = "Epic"
    row[1].text = "EPIC-1"
    row[2].text = "Daily Loan Analysis & Wire Instructions"
    row[3].text = (
        "For each completed daily run, business can identify loans to purchase, "
        "generate wire-ready outputs, and forward data downstream."
    )
    row[4].text = "N/A"
    row[5].text = "None"
    doc.add_paragraph()

    stories = [
        {
            "key": "ST-1.1",
            "summary": "Select loans to purchase from a run",
            "points": 13,
            "depends": "None",
            "sprint": "Sprint 1",
            "ac": [
                "User can view a run-level loans-to-purchase list.",
                "User can select/deselect loans and save final decisions.",
                "Optional notes/reasons for exclusion are captured.",
                "Selections persist and are associated with run + user + timestamp.",
                "Saved decisions reload correctly when revisiting run.",
            ],
            "tasks": [
                ("TSK-1.1-1", "Refine 1.1 acceptance criteria and decision states with Ops/Treasury", 2, "None"),
                ("TSK-1.1-2", "Design DB model: run_purchase_decisions + constraints", 3, "TSK-1.1-1"),
                ("TSK-1.1-3", "Create migration for run_purchase_decisions table", 2, "TSK-1.1-2"),
                ("TSK-1.1-4", "Build backend APIs: list candidates, get decisions, bulk save decisions", 5, "TSK-1.1-3"),
                ("TSK-1.1-5", "Implement UI in Run Detail: select/deselect, reason/note, save", 5, "TSK-1.1-4"),
                ("TSK-1.1-6", "Add audit events for create/update/delete decision actions", 3, "TSK-1.1-4"),
                ("TSK-1.1-7", "Add automated tests (API + UI) and UAT script", 3, "TSK-1.1-5, TSK-1.1-6"),
            ],
        },
        {
            "key": "ST-1.2",
            "summary": "Generate/export wire instruction data",
            "points": 8,
            "depends": "ST-1.1",
            "sprint": "Sprint 2",
            "ac": [
                "User can generate wire export from selected loans only.",
                "Export supports agreed format (CSV/Excel/template).",
                "Export includes required fields: counterparty, amount, reference, loan IDs.",
                "User can download generated file and see generation status.",
            ],
            "tasks": [
                ("TSK-1.2-1", "Finalize wire export field mapping with Treasury", 2, "ST-1.1"),
                ("TSK-1.2-2", "Implement backend wire export service", 3, "TSK-1.2-1"),
                ("TSK-1.2-3", "Add API endpoint for wire export generation/download", 3, "TSK-1.2-2"),
                ("TSK-1.2-4", "Implement frontend generate/download wire export flow", 3, "TSK-1.2-3"),
                ("TSK-1.2-5", "Add tests and sample output fixtures", 2, "TSK-1.2-3"),
            ],
        },
        {
            "key": "ST-1.3",
            "summary": "Wire instruction template and validation",
            "points": 8,
            "depends": "ST-1.2",
            "sprint": "Sprint 3",
            "ac": [
                "Template fields are configurable.",
                "Required wire validations are enforced.",
                "Validation errors are shown clearly to users.",
                "Business rules for amounts/counterparties are documented and traceable.",
            ],
            "tasks": [
                ("TSK-1.3-1", "Define template config schema and storage", 2, "ST-1.2"),
                ("TSK-1.3-2", "Build validation engine for wire fields", 3, "TSK-1.3-1"),
                ("TSK-1.3-3", "Integrate validations into export flow", 2, "TSK-1.3-2"),
                ("TSK-1.3-4", "Add UI messaging for validation failures", 2, "TSK-1.3-3"),
                ("TSK-1.3-5", "Document business rules in ops playbook", 1, "TSK-1.3-2"),
            ],
        },
        {
            "key": "ST-1.4",
            "summary": "Forward loan data for downstream use",
            "points": 8,
            "depends": "ST-1.1",
            "sprint": "Sprint 3",
            "ac": [
                "Selected loans can be exported/sent for downstream owners.",
                "Payload contains required attributes.",
                "Delivery method status is visible (exported/sent/failed).",
                "Format and process are configurable per downstream target.",
            ],
            "tasks": [
                ("TSK-1.4-1", "Define downstream contract with target owners", 2, "ST-1.1"),
                ("TSK-1.4-2", "Implement outbound payload builder", 3, "TSK-1.4-1"),
                ("TSK-1.4-3", "Implement delivery adapter v1 (file export)", 2, "TSK-1.4-2"),
                ("TSK-1.4-4", "Add status tracking + retry metadata", 2, "TSK-1.4-3"),
                ("TSK-1.4-5", "Add integration tests and sample handoff files", 2, "TSK-1.4-3"),
            ],
        },
        {
            "key": "ST-1.5",
            "summary": "Audit trail for daily decisions",
            "points": 5,
            "depends": "ST-1.1, ST-1.2, ST-1.4",
            "sprint": "Sprint 3",
            "ac": [
                "Selections, wire exports, and forwards are logged.",
                "Each event includes user, timestamp, run ID, action, and payload summary.",
                "Audit records are queryable by run and date.",
                "Audit data supports compliance/dispute reviews.",
            ],
            "tasks": [
                ("TSK-1.5-1", "Define audit event taxonomy and schema", 1, "ST-1.1"),
                ("TSK-1.5-2", "Emit audit events from decision, wire, and forward workflows", 3, "TSK-1.5-1"),
                ("TSK-1.5-3", "Create audit query/report endpoint", 2, "TSK-1.5-2"),
                ("TSK-1.5-4", "Add compliance-focused tests", 1, "TSK-1.5-2"),
            ],
        },
        {
            "key": "ST-1.6",
            "summary": "Run-level summary for operations",
            "points": 5,
            "depends": "ST-1.1, ST-1.2",
            "sprint": "Sprint 2",
            "ac": [
                "Run summary shows selected count/amount, exceptions, wire-ready totals.",
                "Wire + forward step statuses visible in one place.",
                "Suitable for daily operations review and handoff.",
            ],
            "tasks": [
                ("TSK-1.6-1", "Define run-summary metrics with Ops", 1, "ST-1.1"),
                ("TSK-1.6-2", "Build backend run summary aggregator", 2, "TSK-1.6-1"),
                ("TSK-1.6-3", "Implement frontend run summary panel/report", 2, "TSK-1.6-2"),
                ("TSK-1.6-4", "Add acceptance tests and ops signoff checklist", 1, "TSK-1.6-3"),
            ],
        },
    ]

    doc.add_heading("Sprint Plan", level=1)
    sprint_table = doc.add_table(rows=1, cols=4)
    h = sprint_table.rows[0].cells
    h[0].text = "Sprint"
    h[1].text = "Stories"
    h[2].text = "Capacity (SP)"
    h[3].text = "Exit Criteria"
    for sprint, story_keys, points, exit_criteria in [
        ("Sprint 1", "ST-1.1", "13", "Selections saved per run; audited; tested; UAT complete."),
        ("Sprint 2", "ST-1.2, ST-1.6", "13", "Wire export downloadable; run-level ops summary available."),
        ("Sprint 3", "ST-1.3, ST-1.4, ST-1.5", "21", "Template/validation, downstream handoff, and full audit trail complete."),
    ]:
        r = sprint_table.add_row().cells
        r[0].text = sprint
        r[1].text = story_keys
        r[2].text = points
        r[3].text = exit_criteria

    for s in stories:
        doc.add_page_break()
        doc.add_heading(f"Story {s['key']}: {s['summary']}", level=1)

        meta = doc.add_table(rows=1, cols=5)
        m = meta.rows[0].cells
        m[0].text = "Issue Type"
        m[1].text = "Story Key"
        m[2].text = "Sprint"
        m[3].text = "Story Points"
        m[4].text = "Dependencies"
        r = meta.add_row().cells
        r[0].text = "Story"
        r[1].text = s["key"]
        r[2].text = s["sprint"]
        r[3].text = str(s["points"])
        r[4].text = s["depends"]

        doc.add_paragraph()
        doc.add_paragraph("Acceptance Criteria", style="Heading 2")
        for ac in s["ac"]:
            doc.add_paragraph(ac, style="List Bullet")

        doc.add_paragraph("Tasks", style="Heading 2")
        t = doc.add_table(rows=1, cols=5)
        th = t.rows[0].cells
        th[0].text = "Issue Type"
        th[1].text = "Task Key"
        th[2].text = "Summary"
        th[3].text = "Points"
        th[4].text = "Depends On"

        for task_key, summary, pts, dep in s["tasks"]:
            rr = t.add_row().cells
            rr[0].text = "Task"
            rr[1].text = task_key
            rr[2].text = summary
            rr[3].text = str(pts)
            rr[4].text = dep

    doc.add_page_break()
    doc.add_heading("Jira Import / Setup Notes", level=1)
    for note in [
        "Create Epic EPIC-1 first, then create stories and link to Epic.",
        "Use issue links: 'blocks' / 'is blocked by' for listed dependencies.",
        "Use custom field 'Story Points' for points values.",
        "Create labels: epic1, loan-selection, wire, downstream, audit, ops-summary.",
        "Attach this document to the Epic for stakeholder alignment.",
    ]:
        doc.add_paragraph(note, style="List Bullet")

    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()

